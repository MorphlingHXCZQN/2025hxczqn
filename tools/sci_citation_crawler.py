"""Utility to fetch highly cited SCI literature via the Crossref API.

The module now supports two complementary workflows:

* Quick tabular exports (Markdown/CSV) for inclusion in research plans.
* Optional full-text collection for the近 5 年内 (last N years) 的 TDP-43 相关
  医学 SCI 文献，包括 PDF/HTML 下载与摘要提取，再写入 Word 报告。

Users can still rely on cached JSON元数据 when offline. When网络可用时，脚本会
尽量下载高分高引用文章的原文；若无法获取，则退回 Crossref 摘要文本。
"""

from __future__ import annotations

import argparse
import dataclasses
import html
import json
import logging
import mimetypes
import re
import sys
from datetime import datetime
from pathlib import Path
from typing import Iterable, List, Optional, Tuple

try:
    import requests  # type: ignore
except ModuleNotFoundError:  # pragma: no cover - optional dependency for offline cache usage
    requests = None

try:  # pragma: no cover - optional dependency
    from docx import Document  # type: ignore
except ModuleNotFoundError:  # pragma: no cover - allow environments without python-docx
    Document = None


CROSSREF_API = "https://api.crossref.org/works"
USER_AGENT = (
    "sci-citation-crawler/1.0 (mailto:research-informatics@example.com)"
)

LOGGER = logging.getLogger(__name__)


@dataclasses.dataclass
class Article:
    """Container for article metadata returned by Crossref."""

    title: str
    doi: str
    year: Optional[int]
    journal: str
    citation_count: int
    authors: str
    abstract: Optional[str] = None
    fulltext_links: List[dict] = dataclasses.field(default_factory=list)


@dataclasses.dataclass
class RetrievalRecord:
    """Stores how content was obtained for an article."""

    mode: str
    note: str = ""
    path: Optional[Path] = None
    abstract_text: Optional[str] = None


def strip_html_tags(value: str) -> str:
    """Remove HTML/XML tags from Crossref abstract strings."""

    if not value:
        return ""
    clean = re.sub(r"<[^>]+>", " ", value)
    return html.unescape(re.sub(r"\s+", " ", clean)).strip()


def sanitise_filename(value: str) -> str:
    """Create a filesystem-friendly filename stem."""

    if not value:
        return "article"
    stem = re.sub(r"[^a-zA-Z0-9._-]", "_", value)
    return stem[:200] or "article"


def guess_extension(content_type: str, url: str) -> str:
    """Infer a sensible file extension for a downloaded asset."""

    content_type = (content_type or "").lower()
    url = url.lower()
    if "pdf" in content_type or url.endswith(".pdf"):
        return ".pdf"
    if "msword" in content_type or url.endswith(".doc"):
        return ".doc"
    if "wordprocessingml" in content_type or url.endswith(".docx"):
        return ".docx"
    if "html" in content_type or url.endswith(".html") or url.endswith(".htm"):
        return ".html"
    guessed = mimetypes.guess_extension(content_type or "")
    return guessed or ".bin"


def get_abstract_text(article: Article) -> Optional[str]:
    """Return a plain-text abstract when available."""

    if not article.abstract:
        return None
    return strip_html_tags(article.abstract)


def download_fulltext(
    article: Article,
    download_dir: Path,
    max_bytes: int = 25_000_000,
    timeout: int = 90,
) -> Tuple[Optional[Path], str]:
    """Attempt to download an article's full text using Crossref links."""

    if requests is None:
        return None, "requests package not available; cannot download full text."
    if not article.fulltext_links:
        return None, "No full-text links advertised for this record."

    download_dir.mkdir(parents=True, exist_ok=True)
    session = requests.Session()
    headers = {"User-Agent": USER_AGENT}
    last_error = ""

    for link in article.fulltext_links:
        url = link.get("URL") or link.get("url")
        if not url:
            continue
        content_type = link.get("content-type", "")
        try:
            response = session.get(url, headers=headers, timeout=timeout)
            response.raise_for_status()
        except requests.RequestException as exc:  # type: ignore[arg-type]
            last_error = f"Failed to fetch {url}: {exc}"
            LOGGER.debug(last_error)
            continue

        content = response.content
        if not content:
            last_error = f"Empty response for {url}"
            LOGGER.debug(last_error)
            continue
        if len(content) > max_bytes:
            last_error = f"Skipped {url} because file size exceeded {max_bytes} bytes"
            LOGGER.debug(last_error)
            continue

        extension = guess_extension(content_type, url)
        stem = sanitise_filename(article.doi or article.title)
        output_path = download_dir / f"{stem}{extension}"
        output_path.write_bytes(content)
        return output_path, ""

    return None, last_error or "No accessible full-text links."


def dump_articles_to_cache(articles: Iterable[Article], cache_path: Path) -> None:
    """Persist article metadata (including abstracts/links) to a cache file."""

    cache_path.parent.mkdir(parents=True, exist_ok=True)
    payload = [
        {
            "title": article.title,
            "doi": article.doi,
            "year": article.year,
            "journal": article.journal,
            "citation_count": article.citation_count,
            "authors": article.authors,
            "abstract": article.abstract,
            "links": article.fulltext_links,
        }
        for article in articles
    ]
    cache_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def write_word_report(
    articles: Iterable[Article],
    records: Iterable[RetrievalRecord],
    output_path: Path,
    window_description: str,
) -> None:
    """Generate a Word document summarising retrieval outcomes."""

    if Document is None:  # pragma: no cover - optional dependency
        raise RuntimeError(
            "The 'python-docx' package is required to produce Word reports. "
            "Install it via 'pip install python-docx' or omit --word-output."
        )

    doc = Document()
    doc.add_heading("TDP-43 Recent Literature Collection", level=1)
    doc.add_paragraph(
        f"Generated on {datetime.now():%Y-%m-%d} for {window_description}."
    )

    for index, (article, record) in enumerate(zip(articles, records), start=1):
        doc.add_heading(f"{index}. {article.title}", level=2)
        doc.add_paragraph(
            f"Journal: {article.journal or 'Unknown'} ({article.year or '未注明'})"
        )
        doc.add_paragraph(f"Citations: {article.citation_count}")
        if article.doi:
            doc.add_paragraph(f"DOI: https://doi.org/{article.doi}")
        if article.authors:
            doc.add_paragraph(f"Authors: {article.authors}")

        doc.add_paragraph(f"Retrieval mode: {record.mode}")
        if record.note:
            doc.add_paragraph(f"Note: {record.note}")
        if record.path:
            doc.add_paragraph(f"Saved file: {record.path}")
        if record.abstract_text:
            doc.add_paragraph("Abstract:")
            doc.add_paragraph(record.abstract_text)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def fetch_articles(
    query: str,
    rows: int = 20,
    from_year: Optional[int] = None,
    issn_filter: Optional[str] = None,
) -> List[Article]:
    """Retrieve highly cited articles from Crossref."""

    if requests is None:
        raise RuntimeError(
            "The 'requests' package is required for live fetching. "
            "Install it or provide a cached dataset via --cache."
        )

    headers = {"User-Agent": USER_AGENT}
    params = {
        "query": query,
        "sort": "is-referenced-by-count",
        "order": "desc",
        "rows": rows,
    }
    filters: List[str] = []
    if from_year is not None:
        filters.append(f"from-pub-date:{from_year}-01-01")
    if issn_filter:
        filters.append(f"issn:{issn_filter}")
    if filters:
        params["filter"] = ",".join(filters)

    response = requests.get(CROSSREF_API, headers=headers, params=params, timeout=30)
    response.raise_for_status()
    items = response.json().get("message", {}).get("items", [])

    articles: List[Article] = []
    for item in items:
        title = item.get("title", ["Unnamed Article"])[0]
        doi = item.get("DOI", "")
        year = None
        if item.get("issued", {}).get("date-parts"):
            year = item["issued"]["date-parts"][0][0]
        journal = item.get("container-title", ["Unknown Journal"])[0]
        citation_count = item.get("is-referenced-by-count", 0)
        authors = ", ".join(
            filter(
                None,
                [
                    " ".join(
                        filter(
                            None,
                            [
                                person.get("given", ""),
                                person.get("family", ""),
                            ],
                        )
                    ).strip()
                    for person in item.get("author", [])
                ],
            )
        )

        articles.append(
            Article(
                title=title,
                doi=doi,
                year=year,
                journal=journal,
                citation_count=citation_count,
                authors=authors,
                abstract=item.get("abstract"),
                fulltext_links=item.get("link", []) or [],
            )
        )

    return articles


def load_cached_articles(cache_path: Path) -> List[Article]:
    """Load article metadata from a JSON cache file."""

    if not cache_path.exists():
        raise FileNotFoundError(f"Cache file not found: {cache_path}")

    raw = json.loads(cache_path.read_text(encoding="utf-8"))

    articles: List[Article] = []
    for item in raw:
        articles.append(
            Article(
                title=item.get("title", "Unnamed Article"),
                doi=item.get("doi", ""),
                year=item.get("year"),
                journal=item.get("journal", "Unknown Journal"),
                citation_count=item.get("citation_count", 0),
                authors=item.get("authors", ""),
                abstract=item.get("abstract"),
                fulltext_links=item.get("links", []) or [],
            )
        )

    return sorted(articles, key=lambda item: item.citation_count, reverse=True)


def to_markdown(articles: Iterable[Article]) -> str:
    """Render a Markdown table from article metadata."""

    header = "| 标题 | 期刊 | 年份 | 引用次数 | DOI | 作者 |\n"
    separator = "| --- | --- | --- | --- | --- | --- |\n"
    rows = []
    for article in articles:
        doi_link = f"https://doi.org/{article.doi}" if article.doi else ""
        rows.append(
            "| {title} | {journal} | {year} | {citations} | {doi} | {authors} |".format(
                title=article.title.replace("|", "／"),
                journal=article.journal.replace("|", "／"),
                year=article.year or "-",
                citations=article.citation_count,
                doi=f"[{article.doi}]({doi_link})" if article.doi else "-",
                authors=(article.authors or "-").replace("|", "／"),
            )
        )
    return header + separator + "\n".join(rows)


def to_csv(articles: Iterable[Article]) -> str:
    """Render a CSV string from article metadata."""

    lines = ["title,journal,year,citation_count,doi,authors"]
    for article in articles:
        line = ",".join(
            json.dumps(value, ensure_ascii=False)
            for value in (
                article.title,
                article.journal,
                article.year,
                article.citation_count,
                article.doi,
                article.authors,
            )
        )
        lines.append(line)
    return "\n".join(lines)


def write_output(content: str, output: Optional[Path]) -> None:
    """Write scraper output to a path or stdout."""

    if output:
        output.parent.mkdir(parents=True, exist_ok=True)
        output.write_text(content, encoding="utf-8")
    else:
        sys.stdout.write(content)


def parse_args(args: Optional[List[str]] = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Fetch highly cited SCI literature for serum TDP-43 studies.",
    )
    parser.add_argument("query", help="Keyword query, e.g. 'serum TDP-43 sepsis'.")
    parser.add_argument(
        "--rows",
        type=int,
        default=20,
        help="Maximum number of articles to fetch (default: 20).",
    )
    parser.add_argument(
        "--from-year",
        type=int,
        help="Lower bound year for publication date filter. Overrides --recent-years.",
    )
    parser.add_argument(
        "--recent-years",
        type=int,
        default=5,
        help=(
            "When --from-year is not supplied, automatically limit the search to the "
            "most recent N years (default: 5)."
        ),
    )
    parser.add_argument(
        "--issn",
        help="Optional ISSN filter to restrict results to a specific journal.",
    )
    parser.add_argument(
        "--format",
        choices=("markdown", "csv"),
        default="markdown",
        help="Output format (default: markdown).",
    )
    parser.add_argument(
        "--output",
        type=Path,
        help="Write results to file (default: stdout).",
    )
    parser.add_argument(
        "--cache",
        type=Path,
        help=(
            "Path to a JSON cache generated from an earlier run. "
            "If provided, data are loaded offline and the query argument is only used "
            "for provenance in logs."
        ),
    )
    parser.add_argument(
        "--save-cache",
        type=Path,
        help="Persist the freshly fetched metadata (including abstracts) to a JSON file.",
    )
    parser.add_argument(
        "--attempt-fulltext",
        action="store_true",
        help="Try downloading full-text assets advertised by Crossref.",
    )
    parser.add_argument(
        "--download-dir",
        type=Path,
        default=Path("outputs/literature/fulltext"),
        help="Directory for saving downloaded PDFs/HTML (default: outputs/literature/fulltext).",
    )
    parser.add_argument(
        "--word-output",
        type=Path,
        help="Optional Word report (*.docx) summarising downloads and abstracts.",
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Skip full-text downloads and Word export while keeping the table output.",
    )
    parser.add_argument(
        "--verbose",
        action="store_true",
        help="Enable verbose logging for debugging download issues.",
    )
    return parser.parse_args(args)


def main(cli_args: Optional[List[str]] = None) -> None:
    options = parse_args(cli_args)
    logging.basicConfig(
        level=logging.DEBUG if options.verbose else logging.INFO,
        format="%(levelname)s:%(name)s:%(message)s",
    )

    current_year = datetime.now().year
    from_year = options.from_year
    if from_year is None:
        recent_years = max(1, options.recent_years)
        from_year = current_year - recent_years + 1
    window_description = f"{from_year}–{current_year}"

    attempt_fulltext = options.attempt_fulltext or bool(options.word_output)

    try:
        if options.cache:
            LOGGER.info("Loading cached articles from %s", options.cache)
            articles = load_cached_articles(options.cache)
            if from_year:
                articles = [
                    item
                    for item in articles
                    if item.year is None or item.year >= from_year
                ]
        else:
            LOGGER.info(
                "Fetching up to %s records for '%s' (from %s)",
                options.rows,
                options.query,
                from_year,
            )
            articles = fetch_articles(
                query=options.query,
                rows=options.rows,
                from_year=from_year,
                issn_filter=options.issn,
            )
            if options.save_cache:
                dump_articles_to_cache(articles, options.save_cache)
                LOGGER.info("Saved cache to %s", options.save_cache)
    except RuntimeError as exc:
        raise SystemExit(str(exc)) from exc
    except requests.RequestException as exc:  # type: ignore[arg-type]
        raise SystemExit(f"Failed to fetch articles: {exc}") from exc

    if not articles:
        raise SystemExit("No articles found for the provided query.")

    articles = articles[: options.rows]
    LOGGER.info("Processing %s articles", len(articles))

    records: List[RetrievalRecord] = []
    for article in articles:
        abstract_text = get_abstract_text(article)

        if options.dry_run:
            mode = "abstract" if abstract_text else "metadata"
            records.append(
                RetrievalRecord(
                    mode=mode,
                    note="Dry-run: skipped downloads.",
                    abstract_text=abstract_text if mode == "abstract" else None,
                )
            )
            continue

        download_note = ""
        if attempt_fulltext:
            download_path, download_error = download_fulltext(
                article, options.download_dir
            )
            if download_path:
                records.append(
                    RetrievalRecord(
                        mode="fulltext",
                        note="Downloaded full text via Crossref link.",
                        path=download_path,
                    )
                )
                continue
            download_note = download_error

        if abstract_text:
            records.append(
                RetrievalRecord(
                    mode="abstract",
                    note=download_note or "Used Crossref abstract.",
                    abstract_text=abstract_text,
                )
            )
        else:
            records.append(
                RetrievalRecord(
                    mode="metadata",
                    note=download_note or "No abstract or accessible full text available.",
                )
            )

    if options.word_output:
        if options.dry_run:
            LOGGER.info(
                "Dry-run enabled; skipping Word export to %s", options.word_output
            )
        else:
            try:
                write_word_report(articles, records, options.word_output, window_description)
                LOGGER.info("Word report saved to %s", options.word_output)
            except RuntimeError as exc:
                raise SystemExit(str(exc)) from exc

    if options.format == "markdown":
        content = to_markdown(articles)
    else:
        content = to_csv(articles)

    write_output(content, options.output)


if __name__ == "__main__":
    main()
